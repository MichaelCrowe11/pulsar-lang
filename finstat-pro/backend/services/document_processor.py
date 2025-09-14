import os
import logging
from typing import Optional, Union
from pathlib import Path
import asyncio
import PyPDF2
import pdfplumber
from docx import Document
import pandas as pd
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for extracting text from various document formats"""
    
    def __init__(self):
        self.max_pages = 500  # Max pages to process
        self.min_text_length = 100  # Minimum text length to consider valid
    
    async def extract_text_from_file(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from various file formats
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif extension == '.txt':
                return await self._extract_from_text(file_path)
            elif extension in ['.doc', '.docx']:
                return await self._extract_from_word(file_path)
            elif extension == '.csv':
                return await self._extract_from_csv(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    async def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber with PyPDF2 fallback"""
        text = ""
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages[:self.max_pages]):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {i+1} ---\n{page_text}"
                        
            if len(text) < self.min_text_length:
                raise ValueError("Insufficient text extracted with pdfplumber")
                
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {str(e)}")
            
            # Fallback to PyPDF2
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = min(len(pdf_reader.pages), self.max_pages)
                
                for i in range(num_pages):
                    page = pdf_reader.pages[i]
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {i+1} ---\n{page_text}"
        
        # Clean up the text
        text = self._clean_text(text)
        
        if len(text) < self.min_text_length:
            raise ValueError(f"Could not extract sufficient text from PDF: {file_path.name}")
        
        return text
    
    async def _extract_from_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            raise
    
    async def _extract_from_word(self, file_path: Path) -> str:
        """Extract text from Word documents"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            
            text = '\n'.join(paragraphs)
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting from Word document: {str(e)}")
            raise
    
    async def _extract_from_csv(self, file_path: Path) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
            
            # Convert DataFrame to readable text format
            text_parts = []
            
            # Add column headers
            text_parts.append("Columns: " + ", ".join(df.columns))
            
            # Add data rows (limit to prevent huge text)
            max_rows = 1000
            for idx, row in df.head(max_rows).iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(row_text)
            
            if len(df) > max_rows:
                text_parts.append(f"\n... and {len(df) - max_rows} more rows")
            
            text = '\n'.join(text_parts)
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Error extracting from CSV: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Fix common OCR issues
        text = text.replace('ﬁ', 'fi')
        text = text.replace('ﬂ', 'fl')
        text = text.replace('™', 'TM')
        text = text.replace('®', '(R)')
        
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    async def extract_financial_sections(self, text: str) -> dict:
        """Extract specific financial sections from document text"""
        sections = {
            'management_discussion': '',
            'financial_statements': '',
            'risk_factors': '',
            'business_overview': '',
            'notes_to_financials': ''
        }
        
        # Define section patterns
        patterns = {
            'management_discussion': r"(?i)(management['']?s?\s+discussion[\s\S]*?(?=\n[A-Z][A-Z\s]{10,}|\Z))",
            'financial_statements': r"(?i)(consolidated\s+financial\s+statements?[\s\S]*?(?=\n[A-Z][A-Z\s]{10,}|\Z))",
            'risk_factors': r"(?i)(risk\s+factors?[\s\S]*?(?=\n[A-Z][A-Z\s]{10,}|\Z))",
            'business_overview': r"(?i)(business\s+overview[\s\S]*?(?=\n[A-Z][A-Z\s]{10,}|\Z))",
            'notes_to_financials': r"(?i)(notes?\s+to\s+.*?financial[\s\S]*?(?=\n[A-Z][A-Z\s]{10,}|\Z))"
        }
        
        for section_name, pattern in patterns.items():
            matches = re.findall(pattern, text)
            if matches:
                sections[section_name] = matches[0][:10000]  # Limit section length
        
        return sections